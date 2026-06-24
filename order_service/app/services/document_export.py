"""
Выгрузка документов в редактируемые форматы:
  - счёт / ТТН / УПД  → XLSX (openpyxl)
  - доверенность / договор → DOCX (python-docx)

Принимает тот же контекст, что и PDF-шаблоны (см. document_service.build_export_ctx
и contract_service). Файлы строятся в памяти и отдаются байтами.
"""
import io

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# ── XLSX-стили ──────────────────────────────────────────────────────────────
_TITLE_FONT = Font(name="Calibri", bold=True, size=13)
_LABEL_FONT = Font(name="Calibri", bold=True, size=10)
_NORMAL     = Font(name="Calibri", size=10)
_HEAD_FONT  = Font(name="Calibri", bold=True, color="FFFFFF", size=10)
_HEAD_FILL  = PatternFill("solid", fgColor="1E3A5F")
_TOTAL_FILL = PatternFill("solid", fgColor="EFEFEF")
_THIN = Side(style="thin", color="BFBFBF")
_BORDER = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)
_CENTER = Alignment(horizontal="center", vertical="center", wrap_text=True)
_LEFT   = Alignment(horizontal="left", vertical="center", wrap_text=True)
_RIGHT  = Alignment(horizontal="right", vertical="center")
_MONEY  = "#,##0.00"


def _g(d, *keys, default=""):
    """Безопасно достать значение из (возможно None) dict."""
    d = d or {}
    for k in keys:
        v = d.get(k)
        if v not in (None, ""):
            return v
    return default


# Excel/LibreOffice трактуют строки, начинающиеся с этих символов, как формулу.
# Клиент с company_name/адресом '=cmd|...' заразит файл, который потом скачает
# менеджер. Префиксуем одинарной кавычкой — она съедается при отображении, но
# рвёт парсер формул. Числа/None не трогаем. (Аналог auth_service._xlsx_safe.)
_XLSX_FORMULA_TRIGGERS = ("=", "+", "-", "@", "\t", "\r")


def _xlsx_safe(value):
    if isinstance(value, str) and value.startswith(_XLSX_FORMULA_TRIGGERS):
        return "'" + value
    return value


def _hrow(ws, row, cols, widths=None):
    for c, title in enumerate(cols, 1):
        cell = ws.cell(row=row, column=c, value=title)
        cell.font, cell.fill, cell.alignment, cell.border = _HEAD_FONT, _HEAD_FILL, _CENTER, _BORDER
    if widths:
        for i, w in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w


def _c(ws, row, col, value, *, bold=False, fmt=None, align=None, fill=None, border=True):
    cell = ws.cell(row=row, column=col, value=_xlsx_safe(value))
    cell.font = Font(name="Calibri", bold=bold, size=10)
    cell.alignment = align or _LEFT
    if border:
        cell.border = _BORDER
    if fmt:
        cell.number_format = fmt
    if fill:
        cell.fill = fill
    return cell


def _party_block(ws, row, seller, buyer, basis=None):
    """Поставщик / Покупатель / (Основание) — слева label, справа значение."""
    def line(r, label, value):
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=1)
        a = ws.cell(row=r, column=1, value=label); a.font = _LABEL_FONT; a.alignment = _LEFT
        ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=6)
        b = ws.cell(row=r, column=2, value=_xlsx_safe(value)); b.font = _NORMAL; b.alignment = _LEFT
    s_inn = _g(seller, "inn")
    s_kpp = _g(seller, "kpp")
    seller_line = f'{_g(seller,"name")}, ИНН {s_inn}' + (f', КПП {s_kpp}' if s_kpp else '') + \
                  (f', {_g(seller,"legal_address")}' if _g(seller, "legal_address") else '')
    b_inn = _g(buyer, "inn")
    buyer_line = _g(buyer, "name") + (f', ИНН {b_inn}' if b_inn and b_inn != "—" else '') + \
                 (f', КПП {_g(buyer,"kpp")}' if _g(buyer, "kpp") else '') + \
                 (f', {_g(buyer,"address","legal_address")}' if _g(buyer, "address", "legal_address") else '')
    line(row, "Поставщик:", seller_line)
    line(row + 1, "Покупатель:", buyer_line)
    nxt = row + 2
    if basis:
        line(nxt, "Основание:", basis); nxt += 1
    return nxt


def _signatures(ws, row, seller):
    director = _xlsx_safe(_g(seller, "director_name"))
    title = _g(seller, "director_title", default="Генеральный директор")
    ws.cell(row=row, column=1, value="Руководитель").font = _LABEL_FONT
    ws.cell(row=row, column=3, value="______________").font = _NORMAL
    ws.cell(row=row, column=5, value=director).font = _NORMAL
    ws.cell(row=row + 1, column=1, value="Гл. бухгалтер").font = _LABEL_FONT
    ws.cell(row=row + 1, column=3, value="______________").font = _NORMAL
    ws.cell(row=row + 1, column=5, value=director).font = _NORMAL
    return row + 2


# ── Счёт ─────────────────────────────────────────────────────────────────────

def invoice_xlsx(ctx: dict) -> bytes:
    wb = Workbook(); ws = wb.active; ws.title = "Счёт"
    ws.sheet_view.showGridLines = False
    ws.merge_cells("A1:F1")
    t = ws["A1"]; t.value = f'Счёт на оплату № {ctx.get("doc_number","")} от {ctx.get("issued_at","")}'
    t.font = _TITLE_FONT; t.alignment = _LEFT
    row = _party_block(ws, 3, ctx.get("seller"), ctx.get("buyer"), ctx.get("basis"))
    row += 1
    _hrow(ws, row, ["№", "Товары (работы, услуги)", "Кол-во", "Ед.", "Цена", "Сумма"],
          widths=[5, 44, 12, 8, 14, 16])
    items = ctx.get("items", [])
    for i, it in enumerate(items, 1):
        r = row + i
        _c(ws, r, 1, i, align=_CENTER)
        _c(ws, r, 2, it.get("name", ""))
        _c(ws, r, 3, float(it.get("qty", 0)), fmt=_MONEY, align=_RIGHT)
        _c(ws, r, 4, it.get("unit", ""), align=_CENTER)
        _c(ws, r, 5, float(it.get("price", 0)), fmt=_MONEY, align=_RIGHT)
        _c(ws, r, 6, float(it.get("sum_no_vat", 0)), fmt=_MONEY, align=_RIGHT)
    row = row + len(items) + 1
    vat_rate = ctx.get("vat_rate")
    for label, val in [("Итого:", ctx.get("subtotal", 0)),
                       (f"НДС {vat_rate}%:" if vat_rate else "Без НДС:", ctx.get("vat_amount", 0)),
                       ("Всего к оплате:", ctx.get("total", 0))]:
        bold = label.startswith("Всего")
        lc = ws.cell(row=row, column=5, value=label); lc.font = Font(bold=bold, size=10); lc.alignment = _RIGHT
        vc = ws.cell(row=row, column=6, value=float(val)); vc.font = Font(bold=bold, size=10)
        vc.number_format = _MONEY; vc.alignment = _RIGHT
        if bold:
            lc.fill = _TOTAL_FILL; vc.fill = _TOTAL_FILL
        row += 1
    row += 1
    w = ws.cell(row=row, column=1, value=f'Сумма прописью: {ctx.get("amount_in_words","")}')
    w.font = Font(bold=True, size=10); ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
    _signatures(ws, row + 2, ctx.get("seller"))
    return _save(wb)


# ── ТТН ──────────────────────────────────────────────────────────────────────

def ttn_xlsx(ctx: dict) -> bytes:
    wb = Workbook(); ws = wb.active; ws.title = "ТТН"
    ws.sheet_view.showGridLines = False
    ws.merge_cells("A1:F1")
    t = ws["A1"]; t.value = f'Товарно-транспортная накладная № {ctx.get("doc_number","")} от {ctx.get("issued_at","")}'
    t.font = _TITLE_FONT; t.alignment = _LEFT
    row = _party_block(ws, 3, ctx.get("seller"), ctx.get("buyer"),
                       basis=f'Заявка № {ctx.get("order_number","")}')
    # доп. строки
    def info(r, label, value):
        a = ws.cell(row=r, column=1, value=label); a.font = _LABEL_FONT
        ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=6)
        b = ws.cell(row=r, column=2, value=_xlsx_safe(value)); b.font = _NORMAL; b.alignment = _LEFT
    info(row, "Адрес доставки:", ctx.get("delivery_address", "")); row += 1
    info(row, "Водитель:", ctx.get("driver_name", "—")); row += 2
    _hrow(ws, row, ["№", "Наименование груза", "Ед.", "Кол-во", "Цена", "Сумма"],
          widths=[5, 44, 8, 14, 14, 16])
    r = row + 1
    _c(ws, r, 1, 1, align=_CENTER)
    _c(ws, r, 2, ctx.get("fuel_name", ""))
    _c(ws, r, 3, "л", align=_CENTER)
    _c(ws, r, 4, float(ctx.get("volume", 0)), fmt=_MONEY, align=_RIGHT)
    _c(ws, r, 5, float(ctx.get("unit_price", 0)), fmt=_MONEY, align=_RIGHT)
    _c(ws, r, 6, float(ctx.get("amount", 0)), fmt=_MONEY, align=_RIGHT)
    row = r + 2
    lc = ws.cell(row=row, column=5, value="Всего:"); lc.font = Font(bold=True, size=10); lc.alignment = _RIGHT
    vc = ws.cell(row=row, column=6, value=float(ctx.get("amount", 0))); vc.font = Font(bold=True, size=10)
    vc.number_format = _MONEY; vc.alignment = _RIGHT; vc.fill = _TOTAL_FILL; lc.fill = _TOTAL_FILL
    row += 2
    w = ws.cell(row=row, column=1, value=f'Сумма прописью: {ctx.get("amount_in_words","")}')
    w.font = Font(bold=True, size=10); ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
    row += 2
    ws.cell(row=row, column=1, value="Отпуск разрешил").font = _LABEL_FONT
    ws.cell(row=row, column=3, value="______________").font = _NORMAL
    ws.cell(row=row, column=5, value=_xlsx_safe(_g(ctx.get("seller"), "director_name"))).font = _NORMAL
    ws.cell(row=row + 1, column=1, value="Груз получил").font = _LABEL_FONT
    ws.cell(row=row + 1, column=3, value="______________").font = _NORMAL
    return _save(wb)


# ── УПД ──────────────────────────────────────────────────────────────────────

def upd_xlsx(ctx: dict) -> bytes:
    wb = Workbook(); ws = wb.active; ws.title = "УПД"
    ws.sheet_view.showGridLines = False
    ws.merge_cells("A1:H1")
    status = "1 — счёт-фактура и передаточный документ" if ctx.get("status_code") == "1" else "2 — передаточный документ"
    t = ws["A1"]; t.value = f'Универсальный передаточный документ № {ctx.get("doc_number","")} от {ctx.get("issued_at","")}'
    t.font = _TITLE_FONT; t.alignment = _LEFT
    st = ws.cell(row=2, column=1, value=f"Статус: {status}"); st.font = Font(italic=True, size=9)
    ws.merge_cells("A2:H2")
    row = _party_block(ws, 4, ctx.get("seller"), ctx.get("buyer"), ctx.get("basis"))
    row += 1
    _hrow(ws, row, ["№", "Наименование", "Кол-во", "Ед.", "Цена",
                    "Сумма без НДС", "НДС", "Сумма с НДС"],
          widths=[5, 34, 10, 8, 13, 15, 13, 15])
    items = ctx.get("items", [])
    for i, it in enumerate(items, 1):
        r = row + i
        _c(ws, r, 1, i, align=_CENTER)
        _c(ws, r, 2, it.get("name", ""))
        _c(ws, r, 3, float(it.get("qty", 0)), fmt=_MONEY, align=_RIGHT)
        _c(ws, r, 4, it.get("unit", ""), align=_CENTER)
        _c(ws, r, 5, float(it.get("price", 0)), fmt=_MONEY, align=_RIGHT)
        _c(ws, r, 6, float(it.get("sum_no_vat", 0)), fmt=_MONEY, align=_RIGHT)
        _c(ws, r, 7, float(it.get("vat", 0)), fmt=_MONEY, align=_RIGHT)
        _c(ws, r, 8, float(it.get("sum", 0)), fmt=_MONEY, align=_RIGHT)
    row = row + len(items) + 1
    vat_rate = ctx.get("vat_rate")
    lc = ws.cell(row=row, column=5, value=f"Итого (НДС {vat_rate}%):" if vat_rate else "Итого (без НДС):")
    lc.font = Font(bold=True, size=10); lc.alignment = _RIGHT
    _c(ws, row, 6, float(ctx.get("subtotal", 0)), bold=True, fmt=_MONEY, align=_RIGHT, fill=_TOTAL_FILL)
    _c(ws, row, 7, float(ctx.get("vat_amount", 0)), bold=True, fmt=_MONEY, align=_RIGHT, fill=_TOTAL_FILL)
    _c(ws, row, 8, float(ctx.get("total", 0)), bold=True, fmt=_MONEY, align=_RIGHT, fill=_TOTAL_FILL)
    row += 2
    director = _xlsx_safe(_g(ctx.get("seller"), "director_name"))
    ws.cell(row=row, column=1, value="Товар передал (руководитель):").font = _LABEL_FONT
    ws.cell(row=row, column=5, value=director).font = _NORMAL
    ws.cell(row=row + 1, column=1, value="Товар получил:").font = _LABEL_FONT
    return _save(wb)


def _save(wb) -> bytes:
    buf = io.BytesIO(); wb.save(buf); return buf.getvalue()


# ── DOCX: доверенность (М-2) ──────────────────────────────────────────────────

def _docx_doc():
    from docx import Document as Docx
    from docx.shared import Pt
    d = Docx()
    style = d.styles["Normal"]
    style.font.name = "Times New Roman"; style.font.size = Pt(11)
    return d


def poa_docx(ctx: dict) -> bytes:
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    d = _docx_doc()
    h = d.add_heading(level=1)
    run = h.add_run(f'ДОВЕРЕННОСТЬ № {ctx.get("doc_number","")}')
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = d.add_paragraph("(на получение товарно-материальных ценностей, форма М-2)")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.font.size = Pt(9); r.italic = True
    d.add_paragraph(f'Дата выдачи: {ctx.get("issued_at","")}     Действительна по: {ctx.get("valid_until","")}')
    buyer = ctx.get("buyer") or {}
    binn = _g(buyer, "inn")
    d.add_paragraph(
        f'Организация-доверитель: {_g(buyer,"name")}'
        + (f', ИНН {binn}' if binn and binn != "—" else '')
        + (f', КПП {_g(buyer,"kpp")}' if _g(buyer, "kpp") else '')
        + (f', {_g(buyer,"address","legal_address")}' if _g(buyer, "address", "legal_address") else '') + '.'
    )
    d.add_paragraph(
        f'Доверенность выдана {ctx.get("driver_name","—")}, паспорт серия '
        f'{ctx.get("passport_series") or "______"} № {ctx.get("passport_number") or "____________"}, '
        f'выдан {ctx.get("passport_issued_by") or "________________________________"} '
        f'{("«"+ctx["passport_issued_at"]+"»") if ctx.get("passport_issued_at") else ""}'.strip() + '.'
    )
    seller = ctx.get("seller") or {}
    d.add_paragraph(
        f'на получение от {_g(seller,"name","short_name")}'
        + (f', ИНН {_g(seller,"inn")}' if _g(seller, "inn") else '')
        + f' товарно-материальных ценностей по заявке № {ctx.get("order_number","")} '
        f'(адрес доставки: {ctx.get("delivery_address","")}).'
    )
    # таблица ТМЦ
    tbl = d.add_table(rows=2, cols=4); tbl.style = "Table Grid"
    for c, v in zip(tbl.rows[0].cells, ["№", "Наименование ТМЦ", "Ед.", "Количество"]):
        c.paragraphs[0].add_run(v).bold = True
    cells = tbl.rows[1].cells
    cells[0].text = "1"; cells[1].text = str(ctx.get("fuel_name", ""))
    cells[2].text = "л"; cells[3].text = f'{ctx.get("volume", 0):,.2f}'.replace(",", " ")
    d.add_paragraph(
        f'Всего на сумму: {ctx.get("amount", 0):,.2f}'.replace(",", " ")
        + f' руб. ({ctx.get("amount_in_words","")}).'
    )
    d.add_paragraph("")
    d.add_paragraph(f'Подпись лица, получившего доверенность: ______________  /{ctx.get("driver_name","")}/')
    d.add_paragraph(
        f'{_g(buyer,"director_title", default="Руководитель")}: ______________  /{_g(buyer,"director_name")}/'
    )
    d.add_paragraph("М.П.")
    buf = io.BytesIO(); d.save(buf); return buf.getvalue()


# ── DOCX: договор поставки ────────────────────────────────────────────────────

def contract_docx(ctx: dict) -> bytes:
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    d = _docx_doc()
    h = d.add_heading(level=1)
    h.add_run(f'ДОГОВОР ПОСТАВКИ НЕФТЕПРОДУКТОВ № {ctx.get("contract_number","")}')
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = d.add_paragraph()
    hdr.add_run(f'г. {ctx.get("city","Санкт-Петербург")}'
                f'\t\t«{ctx.get("signed_day","")}» {ctx.get("signed_month_ru","")} {ctx.get("signed_year","")} года')
    seller = ctx.get("seller") or {}
    buyer = ctx.get("buyer") or {}
    p = d.add_paragraph()
    p.add_run(
        f'{_g(seller,"name")}, именуемое в дальнейшем «Поставщик», в лице '
        f'{ctx.get("seller_title_genitive") or _g(seller,"director_title", default="Генерального директора")} '
        f'{ctx.get("seller_director_genitive") or _g(seller,"director_name", default="________________")}, '
        f'действующего на основании Устава, с одной стороны, и {_g(buyer,"name")}, именуемое в дальнейшем '
        f'«Покупатель», в лице {ctx.get("buyer_title_genitive") or _g(buyer,"director_title", default="руководителя")} '
        f'{ctx.get("buyer_director_genitive") or _g(buyer,"director_name", default="________________")}, '
        f'действующего на основании Устава, с другой стороны, заключили настоящий Договор о нижеследующем.'
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sections = [
        ("1. Предмет договора", [
            "Поставщик обязуется в течение срока действия настоящего Договора поставить, а Покупатель принять и оплатить нефтепродукты, далее именуемые Товар, в порядке и на условиях настоящего Договора.",
            "Наименование, ассортимент, количество, цена, условия поставки оговаривается Сторонами договора в спецификациях, являющихся приложением к договору и его неотъемлемой частью (далее – Спецификация), или на основании выставленного счета. Поставка Товара осуществляется на основании Заявок.",
            "Качество Товара должно соответствовать ГОСТам или Технологическому регламенту и подтверждаться сертификатами или паспортами качества, передаваемыми с каждой поставляемой партией нефтепродукта.",
            "Поставщик гарантирует: соблюдение надлежащих условий хранения Товара до его передачи Покупателю; надлежащее выполнение производственного контроля качества и безопасности, соблюдения требований нормативных и технических документов к условиям изготовления и оборота Товара; наличие документов, сопровождающих оборот Товара.",
        ]),
        ("2. Порядок поставки товара", [
            "Товар по настоящему Договору поставляется по Заявкам партиями в пределах количества и ассортимента, на основании выставленного счета.",
            "При необходимости поставки очередной партии Товара Покупатель оформляет заявку, содержащую указание на количество и ассортимент подлежащего поставке Товара, а также адрес доставки Товара.",
            "Заявки передаются Покупателем Поставщику в письменном виде. Заявка подлежит исполнению Поставщиком в случае, если она была получена до 17 часов дня предшествующего дню осуществления поставки. Поставка Товара, по принятой к исполнению заявке, производится в течение 24 часов с момента принятия заявки Поставщиком.",
            "Поставка каждой партии Товара на склад Покупателя осуществляется за счет Поставщика путем доставки Товара Покупателю по указанному им в заявке адресу. Одновременно с партией Товара Поставщик передает Покупателю следующие документы: универсальный передаточный документ (далее – УПД); паспорт качества или сертификат качества (в случае, если он требуется для данного типа Продукции), выданный заводом-изготовителем; сертификат соответствия (в случае, если он требуется для данного типа Продукции).",
            "Приемка партии Товара по ассортименту, качеству и количеству проводится при передаче Товара Покупателю. Обязательства Поставщика по поставке партии Товара Покупателю считаются выполненными с момента подписания УПД на эту партию Товара представителями Поставщика и Покупателя.",
            "Право собственности на Товар переходит от Поставщика к Покупателю с момента приемки Товара Покупателем и подписания Сторонами УПД.",
            "В случае обнаружения несоответствия поставленного Товара по качеству Покупатель отказывается от приемки Товара и незамедлительно вызывает представителя Поставщика на объект для отбора арбитражной пробы и составления двустороннего акта о несоответствии качества поставленного Товара.",
            "Поставщик не позднее, чем через 2 (два) часа с момента вызова обязан явиться на объект для отбора арбитражной пробы и подписания акта о несоответствии качества поставленного Товара.",
            "Надлежащим подтверждением соответствия качества поставляемого Товара требованиям настоящего Договора является протокол испытаний арбитражной пробы, предоставленной независимой лабораторией, имеющую аккредитацию в Российской Федерации на проведение соответствующих испытаний (выбор лаборатории за Покупателем).",
            "В случае неприбытия представителя Поставщика Покупатель составляет односторонний акт о несоответствии качества поставленного Товара.",
            "Акт является основанием для отказа от приемки Товара, возврата Товара Поставщику и обязанностью Поставщика заменить некачественный Товар в течение 24 (Двадцати четырех) часов с момента составления акта о несоответствии качества поставленного Товара.",
            "В случае поставки Товара в ассортименте, не соответствующем условиям согласованной заявки, Покупатель вправе по своему выбору применить последствия нарушения условий об ассортименте в соответствии со ст. 468 ГК РФ, а Поставщик, в случае отказа от приемки, обязан произвести замену Товара в сроки, предусмотренные настоящим Договором.",
            "Поставщик обязуется еженедельно предоставлять Покупателю УПД.",
        ]),
        ("3. Цены и порядок расчетов", [
            "Товар поставляется по ценам, согласованным Сторонами, на основании выставленного счета.",
            "Цена Товара включает в себя стоимость дополнительных затрат (погрузка, доставка, слив и иные), акциз и НДС. Цена каждой партии Товара указывается в УПД, оформленной на эту партию Товара, а также в выставляемом Поставщиком счете.",
            "Покупатель производит 100% предоплату Товара по ценам, каждой партии в отдельности путем перечисления денежных средств на расчетный счет Поставщика в течение 3-х банковских дней с момента получения счета на оплату.",
            "Днем оплаты считается день поступления денежных средств на расчетный счет Поставщика.",
            "В случае не выборки количества Товара Покупателем в течение срока, согласованного Сторонами в Спецификации, Покупатель имеет право выбрать оставшееся количество не выбранного Товара по тем же ценам в следующий согласованный Сторонами в Спецификации срок, без риска уплаты штрафа или убытков. Покупатель вправе уменьшить или увеличить объем товара.",
        ]),
        ("4. Ответственность сторон", [
            "Любая из Сторон настоящего договора, не исполнившая обязательства по Договору или исполнившая их ненадлежащим образом, несет ответственность при наличии вины (умысла или неосторожности, небрежности, неосмотрительности) в соответствии с действующим законодательством Российской Федерации. Стороны несут ответственность за неисполнение или ненадлежащее исполнение условий настоящего договора. Убытки, включая упущенную выгоду, возникшие у Сторон в результате неисполнения или ненадлежащего исполнения условий настоящего договора, относятся на виновную сторону помимо выплаты ею штрафов и неустойки.",
            "В случае неуплаты платежа Покупателем в срок, указанный в п. 3.3, Покупатель по требованию Поставщика уплачивает Поставщику пени в размере 1% от суммы просроченного платежа за каждый календарный день просрочки.",
            "За нарушение сроков поставки Продукции Поставщик по требованию покупателя уплачивает пени в размере 1% от стоимости непоставленного в срок Товара за каждый календарный день просрочки.",
            "Поставщик несет ответственность за качество доставленного Товара и готов принимать претензии в течение 10 календарных дней с момента поставки.",
            "Стороны освобождаются от ответственности за частичное или полное неисполнение обязательств по настоящему договору, если это неисполнение явилось следствием обстоятельств непреодолимой силы, возникших после заключения договора в результате событий чрезвычайного характера, которые участник не мог ни предвидеть, ни предотвратить разумными мерами (форс-мажор). К таким событиям чрезвычайного характера относятся: наводнение, пожар, землетрясение, взрыв, шторм, оседание почвы, эпидемия и иные явления природы, а также война или военные действия, а также имеющие обязательную силу постановления Правительства Российской Федерации, указы Президента Российской Федерации или распоряжения (указания) иных государственных органов, которые препятствуют исполнению договора.",
            "Заказчик вправе досрочно расторгнуть договор, в любое время уплатив стоимость доставленного топлива.",
        ]),
        ("5. Разрешение споров", [
            "Все возникшие по настоящему Договору споры и разногласия рассматриваются Сторонами путем предъявления письменных претензий. Срок рассмотрения претензии – 10 (десять) календарных дней с момента получения.",
            "Если указанные споры не могут быть решены путем переговоров, они подлежат разрешению в соответствии с действующим законодательством в Арбитражном суде Санкт-Петербурга и Ленинградской области.",
        ]),
        ("6. Заключительные положения", [
            "Любые изменения и дополнения к настоящему договору действительны лишь при условии, если они совершены в письменной форме и подписаны надлежаще уполномоченными на то представителями сторон.",
            "Настоящий договор заключен в 2-х экземплярах, по одному для каждой из сторон. Все приложения к настоящему договору оформляются в письменном виде и являются его неотъемлемой частью.",
            f'Настоящий договор вступает в силу с момента подписания и действует 5 лет{(" (до " + ctx.get("effective_until","") + ")") if ctx.get("effective_until") else ""}, либо до полного исполнения обязательств и произведения расчетов.',
            "Во всем остальном, что не предусмотрено настоящим договором, стороны будут руководствоваться действующим гражданским законодательством.",
        ]),
    ]
    for title, clauses in sections:
        hp = d.add_paragraph(); hp.add_run(title).bold = True
        for i, cl in enumerate(clauses, 1):
            cp = d.add_paragraph(f'{cl}', style="List Number")
            cp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Реквизиты сторон
    rp = d.add_paragraph(); rp.add_run("7. Реквизиты и подписи сторон").bold = True
    tbl = d.add_table(rows=2, cols=2); tbl.style = "Table Grid"
    tbl.rows[0].cells[0].paragraphs[0].add_run("ПОСТАВЩИК").bold = True
    tbl.rows[0].cells[1].paragraphs[0].add_run("ПОКУПАТЕЛЬ").bold = True

    def reqs(party):
        lines = [_g(party, "name")]
        if _g(party, "legal_address"):
            lines.append(f'Адрес: {_g(party,"legal_address")}')
        lines.append(f'ИНН {_g(party,"inn", default="—")}' + (f' / КПП {_g(party,"kpp")}' if _g(party, "kpp") else ''))
        if _g(party, "ogrn"):
            lines.append(f'ОГРН {_g(party,"ogrn")}')
        if _g(party, "bank_name"):
            lines.append(f'Банк: {_g(party,"bank_name")}')
        if _g(party, "bik"):
            lines.append(f'БИК {_g(party,"bik")}')
        if _g(party, "checking_account"):
            lines.append(f'Р/с {_g(party,"checking_account")}')
        if _g(party, "correspondent_account"):
            lines.append(f'К/с {_g(party,"correspondent_account")}')
        return "\n".join(lines)

    tbl.rows[1].cells[0].text = reqs(seller)
    tbl.rows[1].cells[1].text = reqs(buyer)
    d.add_paragraph("")
    sp = d.add_paragraph()
    sp.add_run(f'{_g(seller,"director_title", default="Генеральный директор")}: ______________  /{ctx.get("seller_sign_name","")}/'
               f'\t\t{_g(buyer,"director_title", default="Руководитель")}: ______________  /{ctx.get("buyer_sign_name","")}/')
    buf = io.BytesIO(); d.save(buf); return buf.getvalue()


# ── Диспетчер ─────────────────────────────────────────────────────────────────

def export_document(doc_type: str, ctx: dict) -> tuple[bytes, str, str]:
    """Вернуть (bytes, расширение, mime) для редактируемой выгрузки документа заявки."""
    if doc_type in ("invoice", "invoice_preliminary", "invoice_final"):
        return invoice_xlsx(ctx), "xlsx", XLSX_MIME
    if doc_type == "ttn":
        return ttn_xlsx(ctx), "xlsx", XLSX_MIME
    if doc_type == "upd":
        return upd_xlsx(ctx), "xlsx", XLSX_MIME
    if doc_type == "poa":
        return poa_docx(ctx), "docx", DOCX_MIME
    raise ValueError(f"export not supported for {doc_type}")
